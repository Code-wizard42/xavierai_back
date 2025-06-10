import pymupdf
import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import logging

# Add DOC/DOCX support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    OLE_AVAILABLE = True
except ImportError:
    OLE_AVAILABLE = False

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path):
    text_data = []
    try:
        logger.info(f"Attempting to extract text from PDF: {pdf_path}")

        # Check if file exists
        if not os.path.exists(pdf_path):
            logger.error(f"PDF file does not exist: {pdf_path}")
            return text_data

        # Try using pymupdf
        try:
            with pymupdf.open(pdf_path) as doc:
                logger.info(f"Successfully opened PDF with {len(doc)} pages")
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text()
                    logger.debug(f"Extracted {len(text)} characters from page {page_num}")
                    text_data.append({'page': page_num, 'text': text})
        except ImportError:
            # If pymupdf fails, try using fitz (PyMuPDF's alternative import name)
            logger.warning("pymupdf failed, trying fitz instead")
            import fitz
            with fitz.open(pdf_path) as doc:
                logger.info(f"Successfully opened PDF with fitz, {len(doc)} pages")
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text()
                    logger.debug(f"Extracted {len(text)} characters from page {page_num}")
                    text_data.append({'page': page_num, 'text': text})

        logger.info(f"Successfully extracted text from {len(text_data)} pages")

        # If we got no text but the file exists, add a placeholder
        if not text_data and os.path.exists(pdf_path):
            logger.warning(f"No text extracted from PDF, adding placeholder: {pdf_path}")
            text_data.append({'page': 1, 'text': f"PDF file {os.path.basename(pdf_path)} could not be processed for text extraction."})

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
        # Add a placeholder for failed extraction
        text_data.append({'page': 1, 'text': f"PDF file {os.path.basename(pdf_path)} could not be processed: {str(e)}"})

    return text_data



def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text_data = file.read()
        return text_data
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            text_data = file.read()
        return text_data
    except Exception as e:
        logger.error(f"Error reading text file {file_path}: {e}")
        return ""

def extract_folder_content(folder_path):
    folder_data = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_extension = os.path.splitext(file)[1].lower()

            if file_extension == '.pdf':
                pdf_content = extract_text_from_pdf(file_path)
                folder_data.extend(pdf_content)
            elif file_extension in ['.txt', '.md', '.rst']:
                try:
                    text_content = read_text_file(file_path)
                    relative_path = os.path.relpath(file_path, folder_path)
                    # Format text file content similar to PDF content for consistency
                    folder_data.append({'page': 'file', 'text': text_content, 'path': relative_path, 'filename': file})
                    logger.debug(f"Successfully processed text file: {relative_path}")
                except Exception as e:
                    logger.error(f"Error processing text file {file_path}: {e}")
            elif file_extension == '.docx':
                try:
                    docx_content = extract_text_from_docx(file_path)
                    folder_data.extend(docx_content)
                    logger.debug(f"Successfully processed DOCX file: {file}")
                except Exception as e:
                    logger.error(f"Error processing DOCX file {file_path}: {e}")
            elif file_extension == '.doc':
                try:
                    doc_content = extract_text_from_doc(file_path)
                    folder_data.extend(doc_content)
                    logger.debug(f"Successfully processed DOC file: {file}")
                except Exception as e:
                    logger.error(f"Error processing DOC file {file_path}: {e}")

    logger.info(f"Extracted content from {len(folder_data)} files/pages")
    return folder_data



import logging
from urllib.parse import urlparse, urljoin
import requests
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor
from queue import Queue
import time

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Simple storage for ETag and Last-Modified. In production, use a database.
url_cache = {}

def extract_text_from_url(url, max_depth=2, max_pages=100, timeout=500):
    logger.info(f"Starting extraction from URL: {url} with max depth: {max_depth}, max pages: {max_pages}, timeout: {timeout} seconds")

    start_time = time.time()
    results = []
    visited_urls = set()
    url_queue = Queue()
    url_queue.put((url, 0))  # (url, depth)

    def crawl_page(url, depth):
        if time.time() - start_time > timeout:
            logger.warning(f"Timeout reached. Stopping crawl for URL: {url}")
            return None

        if url in visited_urls:
            logger.debug(f"URL already visited: {url}")
            return None

        if len(visited_urls) >= max_pages:
            logger.warning(f"Max pages limit reached. Stopping crawl for URL: {url}")
            return None

        visited_urls.add(url)
        logger.info(f"Crawling page: {url} at depth: {depth}")

        try:
            headers = {'Cache-Control': 'no-cache', 'Pragma': 'no-cache'}
            if url in url_cache:
                if 'etag' in url_cache[url]:
                    headers['If-None-Match'] = url_cache[url]['etag']
                if 'last_modified' in url_cache[url]:
                    headers['If-Modified-Since'] = url_cache[url]['last_modified']

            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 304:  # Not Modified
                logger.info(f"Content of {url} has not changed since last check.")
                return None  # Return None or cached result

            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            web_data = {
                'url': url,
                'title': soup.title.string if soup.title else '',
                'sections': []
            }

            current_section = None
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
                if element.name in ['h1', 'h2', 'h3']:
                    if current_section:
                        web_data['sections'].append(current_section)
                    current_section = {
                        'heading': element.get_text(strip=True),
                        'content': []
                    }
                elif element.name in ['p', 'li']:
                    if current_section:
                        current_section['content'].append(element.get_text(strip=True))
                    else:
                        current_section = {
                            'heading': 'Introduction',
                            'content': [element.get_text(strip=True)]
                        }

            if current_section:
                web_data['sections'].append(current_section)

            logger.info(f"Extracted {len(web_data['sections'])} sections from {url}")

            # Update ETag and Last-Modified if available
            url_cache[url] = {
                'etag': response.headers.get('ETag'),
                'last_modified': response.headers.get('Last-Modified')
            }

            # Extract links for subpage crawling
            if depth < max_depth:
                base_url = urlparse(url)
                for link in soup.find_all('a', href=True):
                    subpage_url = urljoin(url, link['href'])
                    subpage_parsed = urlparse(subpage_url)
                    if subpage_parsed.netloc == base_url.netloc and subpage_url not in visited_urls:
                        logger.debug(f"Adding subpage to queue: {subpage_url}")
                        url_queue.put((subpage_url, depth + 1))

            return web_data

        except requests.RequestException as e:
            logger.error(f"RequestException error extracting text from URL {url}: {e}")
            return {'url': url, 'title': '', 'sections': [], 'error': f"Error fetching URL: {str(e)}"}
        except Exception as e:
            logger.error(f"Unexpected error processing URL {url}: {e}", exc_info=True)
            return {'url': url, 'title': '', 'sections': [], 'error': f"Unexpected error: {str(e)}"}

    with ThreadPoolExecutor(max_workers=5) as executor:
        while not url_queue.empty() and time.time() - start_time < timeout and len(visited_urls) < max_pages:
            current_url, current_depth = url_queue.get()
            future = executor.submit(crawl_page, current_url, current_depth)
            result = future.result()
            if result:
                results.append(result)

    if not results:
        logger.warning("No content extracted")
        return {'url': url, 'title': '', 'sections': [], 'error': "No content extracted"}

    # Combine results from all crawled pages
    combined_result = results[0]
    for result in results[1:]:
        combined_result['sections'].extend(result['sections'])

    logger.info(f"Successfully extracted text from {len(visited_urls)} pages starting from URL: {url}")

    return combined_result








# import requests
# from bs4 import BeautifulSoup
# from urllib.parse import urlparse, urljoin
# import logging
# from queue import Queue
# from concurrent.futures import ThreadPoolExecutor, as_completed

# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# logger = logging.getLogger(__name__)

# def scrape_website(start_url, min_depth=3, max_pages=50):
#     """
#     Scrape a website starting from the given URL up to specified depth.

#     Args:
#         start_url (str): The starting URL to begin scraping from
#         min_depth (int): Minimum depth to crawl (default: 2)
#         max_pages (int): Maximum number of pages to scrape (default: 50)
#     """
#     visited_urls = set()
#     results = []
#     url_queue = Queue()
#     url_queue.put((start_url, 0))  # (url, current_depth)

#     def process_page(url, depth):
#         if url in visited_urls or len(visited_urls) >= max_pages:
#             return None

#         visited_urls.add(url)
#         logger.info(f"Processing {url} at depth {depth}")

#         try:
#             response = requests.get(url, timeout=10)
#             response.raise_for_status()
#             soup = BeautifulSoup(response.text, 'html.parser')

#             # Extract page content
#             page_data = {
#                 'url': url,
#                 'depth': depth,
#                 'title': soup.title.string if soup.title else 'No title',
#                 'text': ' '.join([p.get_text(strip=True) for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'li'])])
#             }

#             # Find links if we haven't reached max depth
#             if depth < min_depth:
#                 base_url = urlparse(url)
#                 for link in soup.find_all('a', href=True):
#                     next_url = urljoin(url, link['href'])
#                     parsed = urlparse(next_url)

#                     # Only follow links from the same domain
#                     if parsed.netloc == base_url.netloc and next_url not in visited_urls:
#                         url_queue.put((next_url, depth + 1))

#             return page_data

#         except Exception as e:
#             logger.error(f"Error processing {url}: {e}")
#             return None

#     with ThreadPoolExecutor(max_workers=5) as executor:
#         futures = set()

#         while not url_queue.empty() and len(visited_urls) < max_pages:
#             current_url, current_depth = url_queue.get()
#             future = executor.submit(process_page, current_url, current_depth)
#             futures.add(future)

#         for future in as_completed(futures):
#             try:
#                 result = future.result()
#                 if result:
#                     results.append(result)
#             except Exception as e:
#                 logger.error(f"Error getting result: {e}")

#     return results


def extract_text_from_docx(docx_path):
    """Extract text from DOCX files using python-docx"""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available. Install with: pip install python-docx")
        return [{'page': 'file', 'text': f"DOCX file {os.path.basename(docx_path)} could not be processed: python-docx not installed"}]
    
    try:
        logger.info(f"Attempting to extract text from DOCX: {docx_path}")
        
        if not os.path.exists(docx_path):
            logger.error(f"DOCX file does not exist: {docx_path}")
            return []
        
        doc = Document(docx_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        
        return [{'page': 'file', 'text': full_text, 'filename': os.path.basename(docx_path)}]
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        return [{'page': 'file', 'text': f"DOCX file {os.path.basename(docx_path)} could not be processed: {str(e)}"}]


def extract_text_from_doc(doc_path):
    """Extract text from DOC files using basic text extraction"""
    try:
        logger.info(f"Attempting to extract text from DOC: {doc_path}")
        
        if not os.path.exists(doc_path):
            logger.error(f"DOC file does not exist: {doc_path}")
            return []
        
        # Try multiple approaches for DOC files
        text_content = ""
        
        # Method 1: Try using olefile if available
        if OLE_AVAILABLE:
            try:
                import struct
                with open(doc_path, 'rb') as f:
                    # Basic DOC file text extraction (simplified approach)
                    content = f.read()
                    # Look for readable text chunks
                    text_chunks = []
                    i = 0
                    while i < len(content) - 4:
                        try:
                            # Simple heuristic to find text
                            chunk = content[i:i+100]
                            decoded = chunk.decode('utf-8', errors='ignore')
                            if len(decoded) > 10 and decoded.isprintable():
                                cleaned = ''.join(c for c in decoded if c.isprintable() or c.isspace())
                                if len(cleaned.strip()) > 5:
                                    text_chunks.append(cleaned.strip())
                        except:
                            pass
                        i += 50
                    
                    text_content = ' '.join(text_chunks)
                    logger.info(f"Extracted text using olefile method: {len(text_content)} characters")
            except Exception as ole_error:
                logger.warning(f"olefile extraction failed: {ole_error}")
        
        # Method 2: Simple binary text extraction as fallback
        if not text_content:
            try:
                with open(doc_path, 'rb') as f:
                    content = f.read()
                    # Extract readable ASCII/UTF-8 text
                    text_content = content.decode('utf-8', errors='ignore')
                    # Clean up the text
                    import re
                    text_content = re.sub(r'[^\w\s\n\r\t.,!?;:"\'()-]', ' ', text_content)
                    text_content = re.sub(r'\s+', ' ', text_content).strip()
                    logger.info(f"Extracted text using binary method: {len(text_content)} characters")
            except Exception as binary_error:
                logger.error(f"Binary extraction failed: {binary_error}")
                text_content = f"DOC file {os.path.basename(doc_path)} could not be processed with available methods"
        
        if not text_content or len(text_content) < 10:
            text_content = f"DOC file {os.path.basename(doc_path)} was processed but contained little readable text. Consider converting to DOCX for better results."
        
        return [{'page': 'file', 'text': text_content, 'filename': os.path.basename(doc_path)}]
        
    except Exception as e:
        logger.error(f"Error extracting text from DOC: {e}", exc_info=True)
        return [{'page': 'file', 'text': f"DOC file {os.path.basename(doc_path)} could not be processed: {str(e)}"}]
